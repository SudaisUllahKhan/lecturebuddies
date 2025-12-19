import os
import re
import PyPDF2
from docx import Document
from PIL import Image
import pytesseract


# ✅ Configurable Tesseract OCR path with fallbacks
# Priority: ENV var TESSERACT_CMD -> Windows default path -> system PATH
_tesseract_env = os.getenv("TESSERACT_CMD")
_windows_default = r"C:\\Program Files\\Tesseract-OCR\\tesseract.exe"
try:
    if _tesseract_env and os.path.exists(_tesseract_env):
        pytesseract.pytesseract.tesseract_cmd = _tesseract_env
    elif os.name == "nt" and os.path.exists(_windows_default):
        pytesseract.pytesseract.tesseract_cmd = _windows_default
    # else: rely on PATH; if not present, OCR calls will raise which we handle
except Exception:
    # If configuration fails, we let _process_image handle exceptions gracefully
    pass


class DocumentProcessor:
    def __init__(self):
        self.supported_formats = {
            ".txt", ".pdf", ".docx", ".doc", ".png", ".jpg", ".jpeg"
        }

    def process_document(self, filepath):
        """
        Detect file type and extract text safely.
        Always returns a string (even if it's an error).
        """
        try:
            file_extension = os.path.splitext(filepath)[1].lower()

            if file_extension == ".txt":
                return self._process_txt(filepath)
            elif file_extension == ".pdf":
                return self._process_pdf(filepath)
            elif file_extension in [".docx", ".doc"]:
                return self._process_word(filepath)
            elif file_extension in [".png", ".jpg", ".jpeg"]:
                return self._process_image(filepath)
            else:
                return f"[Unsupported file format: {file_extension}]"

        except Exception as e:
            return f"[File processing error: {str(e)}]"

    # ------------------------
    # File Type Processors
    # ------------------------

    def _process_txt(self, filepath):
        """Extract text from .txt files."""
        try:
            with open(filepath, "r", encoding="utf-8") as file:
                content = file.read()
            return self._clean_text(content)
        except UnicodeDecodeError:
            try:
                with open(filepath, "r", encoding="latin-1") as file:
                    content = file.read()
                return self._clean_text(content)
            except Exception as e:
                return f"[Error reading TXT file: {str(e)}]"

    def _process_pdf(self, filepath):
        """Extract text from PDF files."""
        try:
            with open(filepath, "rb") as file:
                pdf_reader = PyPDF2.PdfReader(file)
                content = ""
                for page in pdf_reader.pages:
                    extracted = page.extract_text()
                    if extracted:
                        content += extracted + "\n"
            return self._clean_text(content) if content else "[No text extracted from PDF]"
        except Exception as e:
            return f"[Error reading PDF: {str(e)}]"

    def _process_word(self, filepath):
        """Extract text from Word documents (.docx and .doc)."""
        try:
            doc = Document(filepath)
            content = ""

            # Extract paragraphs
            for paragraph in doc.paragraphs:
                content += paragraph.text + "\n"

            # Extract tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        content += cell.text + " "
                content += "\n"

            return self._clean_text(content) if content else "[No text extracted from Word file]"
        except Exception as e:
            return f"[Error reading Word document: {str(e)}]"

    def _process_image(self, filepath):
        """Extract text from image files using OCR (safe with timeout)."""
        try:
            with Image.open(filepath) as img:
                img = img.convert("RGB")  # ensure format is consistent
                # Add timeout to prevent hanging on large images
                content = pytesseract.image_to_string(img, timeout=30)
                cleaned = self._clean_text(content)
                return cleaned if cleaned else "[No text detected in image]"
        except Exception as e:
            # Be explicit when Tesseract is missing to help users
            if "tesseract is not installed" in str(e).lower() or "not found" in str(e).lower():
                return "[OCR unavailable: Tesseract not found. Install Tesseract or set TESSERACT_CMD]"
            return f"[Error reading image: {str(e)}]"

    # ------------------------
    # Helpers
    # ------------------------

    def _clean_text(self, text):
        """Clean and normalize extracted text."""
        if not text:
            return ""

        # Remove extra whitespace
        text = re.sub(r"\s+", " ", text)

        # Remove unwanted characters but keep punctuation
        text = re.sub(r"[^\w\s\.\,\!\?\;\:\-\(\)\[\]\{\}]", "", text)

        # Normalize spaces
        text = re.sub(r"\s+", " ", text)

        return text.strip()

    def get_document_summary(self, content, max_length=500):
        """Generate a brief summary of the document content."""
        if not content:
            return "[No content available to summarize]"

        if len(content) <= max_length:
            return content

        # Take first few sentences
        sentences = content.split(".")
        summary = ""
        for sentence in sentences:
            if len(summary + sentence) < max_length:
                summary += sentence.strip() + ". "
            else:
                break

        return summary.strip()
