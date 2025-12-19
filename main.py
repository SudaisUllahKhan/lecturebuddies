import streamlit as st
import requests
import PyPDF2
from docx import Document
import io
import os
import json
import time
from dotenv import load_dotenv
from document_processor import DocumentProcessor
import numpy as np
import tempfile
from faster_whisper import WhisperModel
# CHANGED: Replaced crashing local libraries with web-compatible recorder
from streamlit_mic_recorder import mic_recorder

# ==========================
# PAGE CONFIGURATION
# ==========================
st.set_page_config(
    page_title="LectureBuddies - Educational Platform",
    page_icon="🎓",
    layout="wide",
    initial_sidebar_state="expanded"
)

# ==========================
# SESSION STATE INITIALIZATION
# ==========================
def init_session_state():
    """Initialize all session state variables"""
    defaults = {
        # Login/Auth
        "authenticated": False,
        "current_user": None,
        "current_page": "dashboard",

        # Chatbot & Summarization
        "messages": [],
        "uploaded_files": [],
        "document_contents": {},
        "chat_model": "llama-3.1-8b-instant",
        "chat_temperature": 0.7,

        # Quiz Generator
        "quiz_output": None,
        "num_questions": 5,
        "difficulty": "Medium",
        "quiz_model": "llama-3.1-8b-instant",
        "quiz_temperature": 0.7,

        # Live Lecture Recording
        "recording": False,
        "transcript": "",
        "partial_transcript": "",
        "chunks_saved": [],
        "live_transcript": "", # Added this back

        # Dashboard
        "selected_feature": None
    }
    for k, v in defaults.items():
        if k not in st.session_state:
            st.session_state[k] = v

init_session_state()

# ==========================
# LOAD API KEY (UNCHANGED)
# ==========================
load_dotenv()
api_key = os.getenv("GROQ_API_KEY")

# ==========================
# GLOBAL STYLING - LECTUREBUDDIES THEME
# ==========================
st.markdown(
    """
    <style>
    @import url('https://fonts.googleapis.com/css2?family=Poppins:wght@300;400;600;700&display=swap');

    /* Prevent scrolling on login page only */
    .login-page body, .login-page html {
        overflow-y: hidden !important;
        height: 100vh !important;
    }

    /* Main App Styling */
    .stApp {
        font-family: 'Poppins', sans-serif;
        background-color: #f5f7fa;
    }

    /* Hide Streamlit default elements */
    MainMenu {visibility: visible;}
    footer {visibility: visible;}
    header {visibility: visible;}
    
    /* CRITICAL FIX: To ensure Streamlit's default header is completely gone 
       (or transparent), we keep this, as custom headers are removed. 
       If you want the default Streamlit title bar back, delete this block. */
    [data-testid="stHeader"] {
        background: rgba(0,0,0,0) !important; 
        height: 0px !important; 
    }

    /* Main Title */
    .main-title {
        text-align: center;
        font-size: 42px;
        font-weight: 800;
        background: linear-gradient(90deg, #4e54c8, #8f94fb, #4e54c8);
        -webkit-background-clip: text;
        -webkit-text-fill-color: transparent;
        background-clip: text;
        margin-bottom: 8px;
        font-family: 'Poppins', sans-serif;
    }

    /* Tagline */
    .tagline {
        text-align: center;
        font-size: 16px;
        color: #666;
        margin-bottom: 15px;
        font-family: 'Poppins', sans-serif;
    }

    /* Gradient Line */
    hr.gradient {
        border: none;
        height: 3px;
        background: linear-gradient(90deg, #4e54c8, #8f94fb, #4e54c8);
        border-radius: 50px;
        margin: 15px auto;
        max-width: 400px;
    }

    /* ==============================================
       TOP NAVBAR - DELETED (Cleaned)
       ============================================== */
    /* All custom header CSS removed. */

    /* ==============================================
       DASHBOARD CONTENT STYLING FIXES
       ============================================== */

    .main .block-container {
        padding-top: 2rem !important; 
        padding-bottom: 3rem !important;
        padding-left: 3rem !important;
        padding-right: 3rem !important;
    }

    /* Sidebar Fixes: Reduce top padding to lift the menu closer to the top */
    [data-testid="stSidebar"] > div:first-child {
        padding-top: 10px !important; 
        padding-bottom: 10px !important;
    }

    /* ==============================================
       SIDEBAR - Modern WordPress Style (Unchanged)
       ============================================== */
    [data-testid="stSidebar"] {
        background: linear-gradient(180deg, #ffffff 0%, #f8f9ff 100%);
        border-right: 1px solid #e0e7ff;
    }

    [data-testid="stSidebar"] > div:first-child {
        padding-top: 10px !important;
    }

    .sidebar-section {
        padding: 10px 0;
    }

    .sidebar-title {
        font-size: 11px;
        font-weight: 700;
        text-transform: uppercase;
        letter-spacing: 1px;
        color: #9ca3af;
        padding: 0 20px;
        margin-bottom: 12px;
    }

    /* Sidebar Button Styling */
    [data-testid="stSidebar"] button {
        width: 100% !important;
        text-align: left !important;
        padding: 8px 15px !important;
        margin: 1px 0 !important;
        border-radius: 8px !important;
        background: transparent !important;
        border: none !important;
        color: #4b5563 !important;
        font-weight: 500 !important;
        font-size: 12px !important;
        transition: all 0.3s ease !important;
        display: flex !important;
        align-items: center !important;
        gap: 8px !important;
        white-space: nowrap !important;
        min-height: 32px !important;
    }

    [data-testid="stSidebar"] button:hover {
        background: linear-gradient(135deg, #4e54c8, #8f94fb) !important;
        color: white !important;
        transform: translateX(5px);
    }

    /* ==============================================
       LOGIN/SIGNUP PAGES - WordPress Style
       ============================================== */
    .login-container {
        display: flex;
        justify-content: center;
        align-items: center;
        min-height: 80vh;
        padding: 40px 20px;
    }

    .login-box-wrapper {
        background: white;
        padding: 45px 40px;
        border-radius: 16px;
        box-shadow: 0 10px 50px rgba(78, 84, 200, 0.12);
        max-width: 420px;
        width: 100%;
    }

    .login-header {
        text-align: center;
        margin-bottom: 30px;
    }

    .login-title {
        font-size: 32px;
        font-weight: 700;
        color: #4e54c8;
        margin-bottom: 8px;
        font-family: 'Poppins', sans-serif;
    }

    .login-subtitle {
        font-size: 15px;
        color: #667eea;
        font-weight: 400;
        font-family: 'Poppins', sans-serif;
    }

    /* Input Field Styling */
    .login-input-wrapper {
        margin-bottom: 20px;
    }

    .login-label {
        display: block;
        font-size: 14px;
        font-weight: 600;
        color: #374151;
        margin-bottom: 8px;
        font-family: 'Poppins', sans-serif;
    }

    /* Override Streamlit input styling */
    .stTextInput > div > div > input,
    .stTextInput > div > input {
        border: 2px solid #e5e7eb !important;
        border-radius: 10px !important;
        padding: 12px 16px !important;
        font-size: 15px !important;
        font-family: 'Poppins', sans-serif !important;
        transition: all 0.3s ease !important;
        height: auto !important;
        width: 100% !important;
        background: #f9fafb !important;
    }

    .stTextInput > div > div > input:focus,
    .stTextInput > div > input:focus {
        border-color: #667eea !important;
        outline: none !important;
        box-shadow: 0 0 0 4px rgba(102, 126, 234, 0.1) !important;
        background: white !important;
    }

    /* Modern Button Styling */
    .stButton > button {
        width: 100% !important;
        background: linear-gradient(135deg, #667eea 0%, #764ba2 100%) !important;
        color: white !important;
        font-weight: 600 !important;
        font-size: 15px !important;
        padding: 9px 12px !important;
        border: none !important;
        border-radius: 10px !important;
        cursor: pointer !important;
        transition: all 0.3s ease !important;
        font-family: 'Poppins', sans-serif !important;
        box-shadow: 0 4px 15px rgba(102, 126, 234, 0.3) !important;
        margin-top: 10px !important;
    }

    .stButton > button:hover {
        background: linear-gradient(135deg, #764ba2 0%, #667eea 100%) !important;
        transform: translateY(-2px) !important;
        box-shadow: 0 8px 25px rgba(102, 126, 234, 0.4) !important;
    }

    /* Tab Styling */
    .stTabs [data-baseweb="tab-list"] {
        gap: 5px;
        background-color: #f3f4f6;
        border-radius: 12px;
        padding: 6px;
        margin-bottom: 30px;
    }

    .stTabs [data-baseweb="tab"] {
        border-radius: 8px;
        font-family: 'Poppins', sans-serif;
        font-weight: 600;
        font-size: 14px;
        padding: 10px 20px;
        transition: all 0.3s ease;
    }

    .stTabs [aria-selected="true"] {
        background: linear-gradient(135deg, #667eea 0%, #764ba2 100%) !important;
        color: white !important;
    }

    /* Links Styling */
    .login-links {
        text-align: center;
        margin-top: 25px;
        padding-top: 25px;
        border-top: 1px solid #e5e7eb;
        font-family: 'Poppins', sans-serif;
    }

    .login-link {
        color: #667eea;
        text-decoration: none;
        font-size: 14px;
        font-weight: 500;
        transition: color 0.3s ease;
    }

    .login-link:hover {
        color: #764ba2;
        text-decoration: underline;
    }

    /* ==============================================
       DASHBOARD CONTENT STYLING
       ============================================== */
    .dashboard-header {
        background: linear-gradient(135deg, #f8f9ff, #e8ecff);
        padding: 30px;
        border-radius: 16px;
        border: 1px solid #d1d9ff;
        margin-bottom: 30px;
        box-shadow: 0 4px 15px rgba(78, 84, 200, 0.08);
    }

    .welcome-text {
        color: #4e54c8;
        font-size: 28px;
        font-weight: 700;
        margin-bottom: 12px;
        font-family: 'Poppins', sans-serif;
    }

    .welcome-subtitle {
        color: #666;
        font-size: 16px;
        font-family: 'Poppins', sans-serif;
    }

    .feature-card:hover {
        transform: translateY(-4px);
        box-shadow: 0 8px 25px rgba(78, 84, 200, 0.12);
        border-color: #d1d9ff;
    }

    .feature-title {
        color: #4e54c8;
        font-size: 16px;
        font-weight: 600;
        margin-bottom: 8px;
        font-family: 'Poppins', sans-serif;
    }

    .feature-description {
        color: #6b7280;
        font-size: 13px;
        font-family: 'Poppins', sans-serif;
        margin-bottom: 12px;
    }

    /* Chat Styling */
    .user-message {
        text-align: right;
        margin: 10px 0;
    }

    .user-bubble {
        display: inline-block;
        background: linear-gradient(135deg, #4e54c8, #8f94fb);
        color: white;
        padding: 12px 18px;
        border-radius: 18px 18px 4px 18px;
        max-width: 70%;
        font-family: 'Poppins', sans-serif;
        font-size: 14px;
        box-shadow: 0 4px 12px rgba(78, 84, 200, 0.2);
    }

    .assistant-message {
        margin: 10px 0;
        color: #202123;
    }

    .assistant-bubble {
        background: white;
        padding: 12px 18px;
        border-radius: 18px 18px 18px 4px;
        max-width: 70%;
        border: 1px solid #e5e7eb;
        font-family: 'Poppins', sans-serif;
        font-size: 14px;
        box-shadow: 0 2px 8px rgba(0,0,0,0.06);
    }

    /* Quiz Styling */
    .quiz-container {
        background: linear-gradient(135deg, #f8f9ff, #e8ecff);
        padding: 20px;
        border-radius: 16px;
        border: 1px solid #d1d9ff;
        margin-top: 10px;
        box-shadow: 0 4px 15px rgba(78, 84, 200, 0.1);
        max-height: 60vh;
        overflow-y: auto;
    }

    /* Coming Soon Styling */
    .coming-soon {
        text-align: center;
        padding: 80px 40px;
        background: linear-gradient(135deg, #f8f9ff, #e8ecff);
        border-radius: 20px;
        border: 2px dashed #d1d9ff;
    }

    .coming-soon-icon {
        font-size: 72px;
        margin-bottom: 24px;
    }

    .coming-soon-title {
        color: #4e54c8;
        font-size: 28px;
        font-weight: 700;
        margin-bottom: 12px;
        font-family: 'Poppins', sans-serif;
    }

    .coming-soon-text {
        color: #6b7280;
        font-size: 16px;
        font-family: 'Poppins', sans-serif;
    }

    /* Demo Button Styling */
    .demo-button {
        margin-top: 15px;
    }

    /* Input Labels Styling */
    label {
        font-weight: 600 !important;
        font-size: 14px !important;
        color: #374151 !important;
        font-family: 'Poppins', sans-serif !important;
        margin-bottom: 8px !important;
    }

    /* Streamlit Default Element Hides */
    .stDeployButton {
        display: none;
    }

    /* Block Container Styling */
    .block-container {
        padding-top: 30px !important;
    }
    
    /* Remaining boilerplate CSS kept for completeness */
    .sidebar-collapsed-toggle {
        position: fixed;
        left: 0;
        top: 50%;
        transform: translateY(-50%);
        width: 40px;
        height: 40px;
        background: linear-gradient(135deg, #4e54c8, #8f94fb);
        border-radius: 0 20px 20px 0;
        border: none;
        border-right: 2px solid white;
        box-shadow: 2px 2px 10px rgba(0,0,0,0.1);
        cursor: pointer;
        z-index: 999;
        display: flex;
        align-items: center;
        justify-content: center;
        transition: all 0.3s ease;
    }

    .sidebar-collapsed-toggle:hover {
        width: 45px;
        box-shadow: 4px 4px 15px rgba(78, 84, 200, 0.3);
    }

    .sidebar-collapsed-toggle::after {
        content: '→';
        color: white;
        font-size: 18px;
        font-weight: bold;
    }

    /* User Info Hover & Click Effects */
    .user-info {
        cursor: pointer !important;
        transition: all 0.3s ease !important;
    }

    .user-info:hover {
        background: linear-gradient(135deg, #e8ecff, #d6d9ff) !important;
        transform: translateY(-2px);
        box-shadow: 0 4px 12px rgba(78, 84, 200, 0.15) !important;
    }

    /* Profile Dropdown (hidden by default) */
    .profile-dropdown {
        position: absolute;
        top: 50px;
        right: 20px;
        background: white;
        border-radius: 12px;
        box-shadow: 0 8px 25px rgba(0,0,0,0.15);
        min-width: 200px;
        z-index: 1000;
        display: none;
        overflow: hidden;
    }

    .profile-dropdown.show {
        display: block !important;
    }

    .dropdown-item {
        padding: 12px 20px;
        color: #4b5563;
        font-size: 14px;
        font-family: 'Poppins', sans-serif;
        cursor: pointer;
        transition: all 0.3s ease;
    }

    .dropdown-item:hover {
        background: #f8f9ff;
        color: #4e54c8;
    }

    .dropdown-item:first-child {
        padding-top: 16px;
    }

    .dropdown-item:last-child {
        padding-bottom: 16px;
    }
    </style>
    """,
    unsafe_allow_html=True
)

# ==========================
# LOGIN AND SIGNUP SECTION
# ==========================
def show_login_page():
    """Display login/signup page with WordPress-style centered form"""
    # Center column approach for login form
    col1, col2, col3 = st.columns([1, 1.5, 1])
    
    with col2:
        # Main title outside the centered box
        st.markdown("""
            <div style="text-align: center; margin-bottom: 30px;">
                <h1 style="font-size: 42px; font-weight: 800; color: #4e54c8; text-shadow: 2px 2px 4px rgba(78, 84, 200, 0.2); font-family: 'Poppins', sans-serif; margin-bottom: 8px;">
                    🎓 Lecturebuddies
                </h1>
                <p style="color: #666; font-size: 16px; font-family: 'Poppins', sans-serif;">
                    Your intelligent study companion—learn, create, and excel with AI
                </p>
            </div>
        """, unsafe_allow_html=True)
        
        st.markdown(
            """
            <div class="login-header">
                <h2 class="login-title">Welcome Back!</h2>
                <p class="login-subtitle">Sign in to continue to Lecturebuddies</p>
            </div>
            """,
            unsafe_allow_html=True
        )
        
        # Login/Signup Tabs
        tab1, tab2 = st.tabs(["Login", "Sign Up"])

        with tab1:
            username = st.text_input("Username", placeholder="Enter your username", key="login_username")
            password = st.text_input("Password", type="password", placeholder="Enter your password", key="login_password")

            if st.button("Login", key="login_btn", help="Login to your account"):
                if username and password:
                    # Simple authentication (no database)
                    if username == "student" and password == "lecturebuddies":
                        st.session_state.authenticated = True
                        st.session_state.current_user = username
                        st.success("Login successful!")
                        st.rerun()
                    else:
                        st.error("Invalid credentials. Try: username='student', password='lecturebuddies'")
                else:
                    st.warning("Please enter both username and password")
            
            st.markdown("---")
            if st.button("🔑 Try Demo Mode", key="demo_login", help="Quick demo access", use_container_width=True):
                st.session_state.authenticated = True
                st.session_state.current_user = "demo_user"
                st.success("Demo login successful!")
                st.rerun()

        with tab2:
            new_username = st.text_input("Choose Username", placeholder="Enter your username", key="signup_username")
            new_password = st.text_input("Choose Password", type="password", placeholder="Enter your password", key="signup_password")
            confirm_password = st.text_input("Confirm Password", type="password", placeholder="Confirm your password", key="confirm_password")

            if st.button("Create Account", key="signup_btn", help="Create new account"):
                if new_username and new_password and confirm_password:
                    if new_password == confirm_password:
                        st.session_state.authenticated = True
                        st.session_state.current_user = new_username
                        st.success("Account created successfully!")
                        st.rerun()
                    else:
                        st.error("Passwords don't match")
                else:
                    st.warning("Please fill in all fields")

# ==========================
# DASHBOARD LAYOUT SECTION
# ==========================
def show_dashboard():
    """Display main dashboard with sidebar and content area"""
    
    # Initialize profile dropdown state
    if "show_profile_dropdown" not in st.session_state:
        st.session_state.show_profile_dropdown = False

    # Add JavaScript to detect sidebar state and show/hide floating toggle button
    st.markdown(
        """
        <script>
        // Function to check sidebar state and toggle floating button
        function checkSidebarState() {
            const sidebar = document.querySelector('[data-testid="stSidebar"]');
            let floatingToggle = document.getElementById('floating-sidebar-toggle');
            
            if (sidebar) {
                const isCollapsed = sidebar.classList.contains('css-17eq0hr') || 
                                  sidebar.style.transform === 'translateX(-100%)' ||
                                  window.getComputedStyle(sidebar).width === '0px';
                
                if (isCollapsed || !sidebar.offsetParent) {
                    // Sidebar is collapsed - show floating button
                    if (!floatingToggle) {
                        floatingToggle = document.createElement('div');
                        floatingToggle.id = 'floating-sidebar-toggle';
                        floatingToggle.className = 'sidebar-collapsed-toggle';
                        floatingToggle.onclick = function() {
                            // Click the collapse button to expand
                            const toggleBtn = document.querySelector('button[data-testid="baseButton-header"]');
                            if (toggleBtn) toggleBtn.click();
                        };
                        document.body.appendChild(floatingToggle);
                    }
                } else {
                    // Sidebar is expanded - hide floating button
                    if (floatingToggle) {
                        floatingToggle.remove();
                    }
                }
            }
        }
        
        // Check sidebar state on load and periodically
        checkSidebarState();
        setInterval(checkSidebarState, 500);
        
        // Also check when sidebar is toggled
        document.addEventListener('click', function(e) {
            if (e.target.closest('button[data-testid="baseButton-header"]')) {
                setTimeout(checkSidebarState, 350);
            }
        });
        </script>
        """,
        unsafe_allow_html=True
    )

    # --- RESPONSIVE TOP NAVIGATION BAR ---
    with st.container():
        # CSS for horizontal alignment and styling
        st.markdown("""
            <style>
                .nav-container {
                    display: flex;
                    justify-content: space-between;
                    align-items: center;
                    padding: 0.5rem 1rem;
                    background-color: transparent;
                    margin-bottom: 1rem;
                }
                .nav-title {
                    color: #4e54c8;
                    font-family: 'Poppins', sans-serif;
                    font-weight: 700;
                    font-size: 1.8rem;
                    margin: 0;
                    display: flex;
                    align-items: center;
                    gap: 10px;
                }
            </style>
        """, unsafe_allow_html=True)
        
        # Flex-like layout using columns
        col_nav_left, col_nav_spacer, col_nav_right = st.columns([4, 6, 2])
        
        with col_nav_left:
            st.markdown('<h1 class="nav-title">🎓 Lecturebuddies</h1>', unsafe_allow_html=True)
            
        with col_nav_right:
            # Profile Menu logic
            current_user = st.session_state.get("current_user", "Guest")
            user_display = current_user if isinstance(current_user, str) else current_user.get("username", current_user.get("email", "Account"))
            
            with st.popover(f"👤 {user_display}", use_container_width=True):
                st.markdown("### 👤 User Account")
                if isinstance(current_user, dict):
                    st.markdown(f"**Email:** {current_user.get('email', 'N/A')}")
                    st.markdown(f"**Role:** {current_user.get('role', 'User')}")
                else:
                    st.markdown(f"**Status:** {current_user}")
                
                st.markdown("---")
                if st.button("🚪 Log Out", key="nav_logout_btn", use_container_width=True, type="primary"):
                    st.session_state.authenticated = False
                    st.session_state.current_user = None
                    st.session_state.selected_feature = None
                    st.rerun()

    st.markdown('<hr class="gradient" style="margin-top: 0; margin-bottom: 2rem; max-width: 100%;">', unsafe_allow_html=True)
    
    # Create Layout - use st.sidebar for proper sidebar integration
    # Features in Sidebar
    with st.sidebar:
        # st.markdown("## 🎓 Main Menu") removed in favor of expander title
        
        # Determine if expander should be open (open if no feature selected)
        is_menu_expanded = st.session_state.selected_feature is None
        
        with st.expander("🎓 Main Menu", expanded=is_menu_expanded):
            # Navigation menu
            if st.button("🏠 Dashboard Home", key="nav_dashboard", use_container_width=True):
                st.session_state.selected_feature = None
                st.rerun()
            
            if st.button("💬 Chatbot & Summarization", key="nav_chatbot", use_container_width=True):
                st.session_state.selected_feature = "chatbot"
                st.rerun()
            
            if st.button("📝 Quiz Generator", key="nav_quiz", use_container_width=True):
                st.session_state.selected_feature = "quiz"
                st.rerun()
            
            if st.button("🎤 Live Lecture Recording", key="nav_recording", use_container_width=True):
                st.session_state.selected_feature = "recording"
                st.rerun()
            
            if st.button("📖 Flash Cards", key="nav_flashcards", use_container_width=True):
                st.session_state.selected_feature = "flashcards"
                st.rerun()
            
            if st.button("🌐 Translation", key="nav_translation", use_container_width=True):
                st.session_state.selected_feature = "translation"
                st.rerun()
            
            if st.button("📋 Notes Manager", key="nav_notes", use_container_width=True):
                st.session_state.selected_feature = "notes"
                st.rerun()
            
            if st.button("👨‍💼 Admin Dashboard", key="nav_admin", use_container_width=True):
                st.session_state.selected_feature = "admin"
                st.rerun()
            
            if st.button("🔍 Search", key="nav_search", use_container_width=True):
                st.session_state.selected_feature = "search"
                st.rerun()
            
            if st.button("📱 Offline Mode", key="nav_offline", use_container_width=True):
                st.session_state.selected_feature = "offline"
                st.rerun()
            
            if st.button("🚪 Logout", key="logout_btn", help="Logout from your account", use_container_width=True):
                st.session_state.authenticated = False
                st.session_state.current_user = None
                st.session_state.selected_feature = None
                st.rerun()

    # Main Content Area
    if st.session_state.selected_feature is None:
        show_welcome_screen()
    elif st.session_state.selected_feature == "chatbot":
        show_chatbot_feature()
    elif st.session_state.selected_feature == "quiz":
        show_quiz_feature()
    elif st.session_state.selected_feature == "recording":
        show_recording_feature()
    elif st.session_state.selected_feature == "flashcards":
        show_flashcards_feature()
    elif st.session_state.selected_feature == "translation":
        show_translation_feature()
    elif st.session_state.selected_feature == "notes":
        show_notes_feature()
    elif st.session_state.selected_feature == "admin":
        show_admin_feature()
    elif st.session_state.selected_feature == "search":
        show_search_feature()
    elif st.session_state.selected_feature == "offline":
        show_offline_feature()
    else:
        show_coming_soon_feature(st.session_state.selected_feature)

# ==========================
# WELCOME SCREEN SECTION
# ==========================
def show_welcome_screen():
    """Display welcome screen with learning visuals"""
    # WELCOME MESSAGE CARD (THEMED BG)
    st.markdown(
        """
        <div class="dashboard-header" style="
            background: linear-gradient(135deg, #e8ecff, #f8f9ff); /* Lightened theme gradient */
            border: 1px solid #d1d9ff; 
            padding: 30px;
            border-radius: 16px;
            margin-bottom: 30px;
            box-shadow: 0 4px 15px rgba(78, 84, 200, 0.08);
        ">
            <h2 class="welcome-text">Welcome to Lecturebuddies Dashboard!</h2>
            <p style="color: #4e54c8; font-size: 16px; font-family: 'Poppins', sans-serif; font-weight: 500;">
                Your comprehensive educational platform powered by AI. Choose a feature from the sidebar to get started.
            </p>
        </div>
        """,
        unsafe_allow_html=True
    )

    # Learning Statistics Cards
    col1, col2, col3, col4 = st.columns(4)

    # --- CARD 1: Study Sessions (Themed BG) ---
    with col1:
        st.markdown(
            """
            <div class="feature-card" style="
                background: linear-gradient(135deg, #f8f9ff 0%, #e8ecff 100%);
                padding: 24px;
                border-radius: 16px;
                border: 1px solid #d1d9ff;
                box-shadow: 0 2px 10px rgba(0,0,0,0.04);
            ">
                <h3 class="feature-title">Study Sessions</h3>
                <p class="feature-description">Track your learning progress</p>
                <h2 style="color: #4e54c8; font-size: 32px; margin: 10px 0;">24</h2>
            </div>
            """,
            unsafe_allow_html=True
        )

    # --- CARD 2: Quizzes Created (Themed BG) ---
    with col2:
        st.markdown(
            """
            <div class="feature-card" style="
                background: linear-gradient(135deg, #f8f9ff 0%, #e8ecff 100%);
                padding: 24px;
                border-radius: 16px;
                border: 1px solid #d1d9ff;
                box-shadow: 0 2px 10px rgba(0,0,0,0.04);
            ">
                <h3 class="feature-title">Quizzes Created</h3>
                <p class="feature-description">Interactive learning materials</p>
                <h2 style="color: #4e54c8; font-size: 32px; margin: 10px 0;">12</h2>
            </div>
            """,
            unsafe_allow_html=True
        )

    # --- CARD 3: Recordings (Themed BG) ---
    with col3:
        st.markdown(
            """
            <div class="feature-card" style="
                background: linear-gradient(135deg, #f8f9ff 0%, #e8ecff 100%);
                padding: 24px;
                border-radius: 16px;
                border: 1px solid #d1d9ff;
                box-shadow: 0 2px 10px rgba(0,0,0,0.04);
            ">
                <h3 class="feature-title">Recordings</h3>
                <p class="feature-description">Audio content processed</p>
                <h2 style="color: #4e54c8; font-size: 32px; margin: 10px 0;">8</h2>
            </div>
            """,
            unsafe_allow_html=True
        )

    # --- CARD 4: Progress (Themed BG) ---
    with col4:
        st.markdown(
            """
            <div class="feature-card" style="
                background: linear-gradient(135deg, #f8f9ff 0%, #e8ecff 100%);
                padding: 24px;
                border-radius: 16px;
                border: 1px solid #d1d9ff;
                box-shadow: 0 2px 10px rgba(0,0,0,0.04);
            ">
                <h3 class="feature-title">Progress</h3>
                <p class="feature-description">Learning efficiency</p>
                <h2 style="color: #4e54c8; font-size: 32px; margin: 10px 0;">85%</h2>
            </div>
            """,
            unsafe_allow_html=True
        )

    # Inspirational Quote (Themed BG)
    st.markdown(
        """
        <div class="feature-card" style="
            text-align: center; 
            margin-top: 30px;
            background: linear-gradient(135deg, #f8f9ff 0%, #e8ecff 100%); /* Light Theme Gradient BG */
            padding: 30px;
            border-radius: 16px;
            border: 1px solid #d1d9ff;
            box-shadow: 0 4px 15px rgba(78, 84, 200, 0.08);
        ">
            <h3 style="
                color: #4e54c8; 
                font-size: 20px; 
                font-weight: 700; 
                margin-bottom: 15px; 
                font-family: 'Poppins', sans-serif;
            ">
                Today's Learning Quote
            </h3>
            <p style="
                color: #374151; 
                font-size: 18px; 
                font-style: italic; 
                font-family: 'Poppins', sans-serif; 
                margin: 0 0 10px 0;
            ">
                "The capacity to learn is a gift; the ability to learn is a skill; the willingness to learn is a choice."
            </p>
            <p style="
                color: #6b7280; 
                font-size: 14px; 
                margin-top: 10px; 
                font-family: 'Poppins', sans-serif;
            ">
                — Brian Herbert
            </p>
        </div>
        """,
        unsafe_allow_html=True
    )

# ==========================
# CHATBOT AND SUMMARIZATION SECTION
# ==========================
def show_chatbot_feature():
    import streamlit as st
    import requests
    import PyPDF2
    from docx import Document
    import io
    import os
    from dotenv import load_dotenv
    import os
    import re
    import PyPDF2
    from docx import Document
    from PIL import Image
    import pytesseract
    
    
    import os
    import re
    import PyPDF2
    from docx import Document
    from PIL import Image
    import pytesseract
    
    
    # ✅ Configurable Tesseract OCR path with fallbacks
    # Priority: ENV var TESSERACT_CMD -> Windows default path -> system PATH
    _tesseract_env = os.getenv("TESSERACT_CMD")
    _windows_default = r"C:\\Program Files\\Tesseract-OCR\\tesseract.exe"
    try:
        if _tesseract_env and os.path.exists(_tesseract_env):
            pytesseract.pytesseract.tesseract_cmd = _tesseract_env
        elif os.name == "nt" and os.path.exists(_windows_default):
            pytesseract.pytesseract.tesseract_cmd = _windows_default
        # else: rely on PATH; if not present, OCR calls will raise which we handle
    except Exception:
        # If configuration fails, we let _process_image handle exceptions gracefully
        pass
    
    
    class DocumentProcessor:
        def __init__(self):
            self.supported_formats = {
                ".txt", ".pdf", ".docx", ".doc", ".png", ".jpg", ".jpeg"
            }
    
        def process_document(self, filepath):
            """
            Detect file type and extract text safely.
            Always returns a string (even if it's an error).
            """
            try:
                file_extension = os.path.splitext(filepath)[1].lower()
    
                if file_extension == ".txt":
                    return self._process_txt(filepath)
                elif file_extension == ".pdf":
                    return self._process_pdf(filepath)
                elif file_extension in [".docx", ".doc"]:
                    return self._process_word(filepath)
                elif file_extension in [".png", ".jpg", ".jpeg"]:
                    return self._process_image(filepath)
                else:
                    return f"[Unsupported file format: {file_extension}]"
    
            except Exception as e:
                return f"[File processing error: {str(e)}]"
    
        # ------------------------
        # File Type Processors
        # ------------------------
    
        def _process_txt(self, filepath):
            """Extract text from .txt files."""
            try:
                with open(filepath, "r", encoding="utf-8") as file:
                    content = file.read()
                return self._clean_text(content)
            except UnicodeDecodeError:
                try:
                    with open(filepath, "r", encoding="latin-1") as file:
                        content = file.read()
                    return self._clean_text(content)
                except Exception as e:
                    return f"[Error reading TXT file: {str(e)}]"
    
        def _process_pdf(self, filepath):
            """Extract text from PDF files."""
            try:
                with open(filepath, "rb") as file:
                    pdf_reader = PyPDF2.PdfReader(file)
                    content = ""
                    for page in pdf_reader.pages:
                        extracted = page.extract_text()
                        if extracted:
                            content += extracted + "\n"
                return self._clean_text(content) if content else "[No text extracted from PDF]"
            except Exception as e:
                return f"[Error reading PDF: {str(e)}]"
    
        def _process_word(self, filepath):
            """Extract text from Word documents (.docx and .doc)."""
            try:
                doc = Document(filepath)
                content = ""
    
                # Extract paragraphs
                for paragraph in doc.paragraphs:
                    content += paragraph.text + "\n"
    
                # Extract tables
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            content += cell.text + " "
                    content += "\n"
    
                return self._clean_text(content) if content else "[No text extracted from Word file]"
            except Exception as e:
                return f"[Error reading Word document: {str(e)}]"
    
        def _process_image(self, filepath):
            """Extract text from image files using OCR (safe with timeout)."""
            try:
                with Image.open(filepath) as img:
                    img = img.convert("RGB")  # ensure format is consistent
                    # Add timeout to prevent hanging on large images
                    content = pytesseract.image_to_string(img, timeout=30)
                    cleaned = self._clean_text(content)
                    return cleaned if cleaned else "[No text detected in image]"
            except Exception as e:
                # Be explicit when Tesseract is missing to help users
                if "tesseract is not installed" in str(e).lower() or "not found" in str(e).lower():
                    return "[OCR unavailable: Tesseract not found. Install Tesseract or set TESSERACT_CMD]"
                return f"[Error reading image: {str(e)}]"
    
        # ------------------------
        # Helpers
        # ------------------------
    
        def _clean_text(self, text):
            """Clean and normalize extracted text."""
            if not text:
                return ""
    
            # Remove extra whitespace
            text = re.sub(r"\s+", " ", text)
    
            # Remove unwanted characters but keep punctuation
            text = re.sub(r"[^\w\s\.\,\!\?\;\:\-\(\)\[\]\{\}]", "", text)
    
            # Normalize spaces
            text = re.sub(r"\s+", " ", text)
    
            return text.strip()
    
        def get_document_summary(self, content, max_length=500):
            """Generate a brief summary of the document content."""
            if not content:
                return "[No content available to summarize]"
    
            if len(content) <= max_length:
                return content
    
            # Take first few sentences
            sentences = content.split(".")
            summary = ""
            for sentence in sentences:
                if len(summary + sentence) < max_length:
                    summary += sentence.strip() + ". "
                else:
                    break
    
            return summary.strip()
    
    
    
    
    
    
    doc_processor = DocumentProcessor()
    
    
    
    # ---------------------------
    # Page Config
    # ---------------------------
    # st.set_page_config(page_title="LectureBuddies - AI Chatbot", page_icon="🤖", layout="wide", initial_sidebar_state="expanded")
    
    # ---------------------------
    # Load API Key from .env
    # ---------------------------
    # (Already loaded globally via st.secrets check)
    
    if not api_key:
        st.error("⚠️ API key missing! Please check your .env file.")
        st.stop()
    
    # ---------------------------
    # Session Initialization
    # ---------------------------
    def init_session_state():
        defaults = {
            "messages": [],
            "uploaded_files": [],
            "document_contents": {}
        }
        for k, v in defaults.items():
            if k not in st.session_state:
                st.session_state[k] = v
    
    init_session_state()
    
    # ---------------------------
    # Inject File Content Helper
    # ---------------------------
    def inject_file_content(user_message: str) -> str:
        """
        Replace file references in user message with extracted text,
        so model never says 'I can't see images'.
        """
        for fname, content in st.session_state.document_contents.items():
            if fname.lower() in user_message.lower():
                extracted = content if content.strip() else "[No text extracted from this file]"
                user_message = user_message.replace(
                    fname,
                    f"(Extracted content: {extracted[:1000]}...)"
                )
        return user_message
    
    # ---------------------------
    # API Interaction
    # ---------------------------
    def get_groq_response(user_input, model="llama-3.1-8b-instant"):
    
        """Send query + context to Groq API and return assistant response."""
        if not api_key or api_key.strip() == "":
            return "⚠️ Missing API key. Please set GROQ_API_KEY in your .env file."
    
        url = "https://api.groq.com/openai/v1/chat/completions"
        headers = {"Authorization": f"Bearer {api_key}", "Content-Type": "application/json"}
    
        # Build document context
        doc_context = ""
        if st.session_state.document_contents:
            doc_context = "\n\n**Available Documents:**\n"
            for fname, content in st.session_state.document_contents.items():
                snippet = content[:1000] if content else "[No extractable text]"
                doc_context += f"\n--- {fname} ---\n{snippet}...\n"
    
        # Build conversation
        messages = [{"role": m["role"], "content": m["content"]} for m in st.session_state.messages]
        system_msg = (
            f"You are LectureBuddies, an AI chatbot designed for education. "
            f"Answer clearly, summarize effectively, and explain concepts step by step."
            f"{' Available Documents: ' + doc_context if doc_context else ''}\n\n"
            "Guidelines:\n"
            "📚 Education-focused\n"
            "📝 Summarization expert\n"
            "🎯 Clarity first (simple language, then details)\n"
            "✅ Confidence + accuracy\n"
            "Break down topics step-by-step, use examples, and stay professional yet supportive."
        )
        messages.insert(0, {"role": "system", "content": system_msg})
        messages.append({"role": "user", "content": user_input})
    
        payload = {"model": model, "messages": messages, "temperature": 0.7, "max_tokens": 1000}
    
        try:
            resp = requests.post(url, headers=headers, json=payload, timeout=30)
            if resp.status_code == 200:
                data = resp.json()
                return data.get("choices", [{}])[0].get("message", {}).get("content", "⚠️ No response received.")
            elif resp.status_code == 401:
                return "❌ Invalid API key. Please check GROQ_API_KEY in your .env file."
            elif resp.status_code == 429:
                return "⏳ Too many requests. Please slow down and retry shortly."
            else:
                return f"⚠️ API Error {resp.status_code}: {resp.text}"
        except requests.exceptions.Timeout:
            return "⏳ Request timed out. Please retry."
        except requests.exceptions.RequestException as e:
            return f"🌐 Network error: {e}"
        except Exception as e:
            return f"⚠️ Unexpected error: {e}"
    
    # ---------------------------
    # Document Processing
    # ---------------------------
    def process_document(uploaded_file):
        """Extract text from uploaded documents (txt, pdf, docx, images with OCR)."""
        try:
            # Save the uploaded file temporarily
            temp_path = os.path.join("temp", uploaded_file.name)
            os.makedirs("temp", exist_ok=True)
            with open(temp_path, "wb") as f:
                f.write(uploaded_file.read())
    
            # Use your DocumentProcessor
            content = doc_processor.process_document(temp_path)
    
            return content if content.strip() else "[No text extracted]"
        except Exception as e:
            return f"[File processing error: {e}]"
    
    # ---------------------------
    # Enhanced Styling (Matching Quiz Generator Theme)
    # ---------------------------
    st.markdown(
        """
        <style>
        @import url('https://fonts.googleapis.com/css2?family=Poppins:wght@300;400;600;700&display=swap');
        
        .main-title {
            text-align: center;
            font-size: 40px;
            font-weight: 800;
            background: linear-gradient(90deg, #4e54c8, #8f94fb, #4e54c8);
            -webkit-background-clip: text;
            -webkit-text-fill-color: transparent;
            background-clip: text;
            margin-bottom: -5px;
            font-family: 'Poppins', sans-serif;
        }
        .tagline {
            text-align: center;
            font-size: 16px;
            color: #666;
            margin-bottom: 15px;
            font-family: 'Poppins', sans-serif;
        }
        hr.gradient {
            border: none;
            height: 3px;
            background: linear-gradient(90deg, #4e54c8, #8f94fb, #4e54c8);
            border-radius: 50px;
            margin: 15px 0;
        }
        .custom-btn {
            display: inline-flex;
            align-items: center;
            gap: 6px;
            padding: 10px 24px;
            font-size: 14px;
            font-weight: 700;
            color: white;
            background: linear-gradient(90deg, #4e54c8, #8f94fb);
            border: none;
            border-radius: 25px;
            cursor: pointer;
            transition: all 0.3s ease;
            text-align: center;
            font-family: 'Poppins', sans-serif;
            box-shadow: 0 3px 10px rgba(78, 84, 200, 0.3);
            margin: 5px 0;
        }
        .custom-btn:hover {
            background: linear-gradient(90deg, #8f94fb, #4e54c8);
            transform: scale(1.03) translateY(-1px);
            box-shadow: 0 5px 15px rgba(78, 84, 200, 0.4);
        }
        .clear-btn {
            background: linear-gradient(90deg, #ff6b6b, #ee5a52);
            color: white;
            border: none;
            border-radius: 15px;
            padding: 8px 16px;
            font-weight: 600;
            cursor: pointer;
            transition: all 0.3s ease;
            font-size: 14px;
            margin: 5px 0;
        }
        .clear-btn:hover {
            background: linear-gradient(90deg, #ee5a52, #ff6b6b);
            transform: scale(1.03);
        }
        .chat-header {
            color: #4e54c8;
            font-size: 22px;
            font-weight: 700;
            margin-bottom: 10px;
            text-align: center;
            font-family: 'Poppins', sans-serif;
        }
        .user-message {
            text-align: right;
            margin: 6px 0;
        }
        .user-bubble {
            display: inline-block;
            background: linear-gradient(90deg, #4e54c8, #8f94fb);
            color: white;
            padding: 10px 14px;
            border-radius: 18px 18px 4px 18px;
            max-width: 70%;
            font-family: 'Poppins', sans-serif;
            font-size: 14px;
            box-shadow: 0 2px 8px rgba(78, 84, 200, 0.3);
        }
        .assistant-message {
            margin: 6px 0;
            color: #202123;
        }
        .assistant-bubble {
            background: white;
            padding: 10px 14px;
            border-radius: 18px 18px 18px 4px;
            max-width: 70%;
            border: 1px solid #e0e0e0;
            font-family: 'Poppins', sans-serif;
            font-size: 14px;
            box-shadow: 0 2px 8px rgba(0,0,0,0.05);
        }
        .welcome-container {
            text-align: center;
            padding: 20px;
            background: linear-gradient(135deg, #f8f9ff, #e8ecff);
            border-radius: 15px;
            border: 1px solid #d1d9ff;
            margin: 10px 0;
        }
        .welcome-title {
            color: #4e54c8;
            font-size: 20px;
            font-weight: 600;
            margin-bottom: 10px;
            font-family: 'Poppins', sans-serif;
        }
        .welcome-text {
            color: #666;
            font-size: 14px;
            font-family: 'Poppins', sans-serif;
        }
        /* Professional chat input styling */
        .stChatInput > div > div > input {
            border: 2px solid #e0e0e0;
            border-radius: 25px;
            padding: 12px 16px;
            font-family: 'Poppins', sans-serif;
            font-size: 14px;
            background: white;
            box-shadow: 0 2px 8px rgba(0,0,0,0.05);
            transition: all 0.3s ease;
        }
        .stChatInput > div > div > input:focus {
            border-color: #4e54c8;
            box-shadow: 0 4px 12px rgba(78, 84, 200, 0.2);
            outline: none;
        }
        /* Fixed page height - prevent scrolling */
        .stApp {
            max-height: 100vh;
            overflow: hidden;
        }
        .block-container {
            padding-top: 1rem;
            padding-bottom: 0rem;
            padding-left: 1rem;
            padding-right: 1rem;
            max-height: 90vh;
            overflow-y: auto;
        }
        /* Reduce sidebar padding */
        .css-1d391kg {
            padding: 0.5rem;
        }
        /* Style sidebar elements */
        .stSidebar > div > div {
            font-family: 'Poppins', sans-serif;
        }
        </style>
        """,
        unsafe_allow_html=True
    )
    
    # ---------------------------
    # Sidebar (Enhanced to Match Theme)
    # ---------------------------
    with st.sidebar:
        st.markdown("## ⚙️ Chat Settings")
        st.markdown("### Customize your experience")
        
        if st.button("🗑️ Clear Chat", key="clear_chat", help="Start a new conversation"):
            st.session_state.messages.clear()
            st.rerun()
        
        st.markdown("---")
        st.markdown("**📎 Upload Files**")
        sidebar_upload = st.file_uploader(
            "Choose files",
            type=['txt', 'pdf', 'docx', 'doc', 'png', 'jpg', 'jpeg', 'gif', 'bmp'],
            key="sidebar_uploader",
            label_visibility="collapsed",
            help="Upload documents or images for context (PDF, DOCX, TXT, Images with OCR)"
        )
        if sidebar_upload and sidebar_upload.name not in st.session_state.document_contents:
            file_details = {
                "filename": sidebar_upload.name,
                "filetype": sidebar_upload.type,
                "filesize": sidebar_upload.size
            }
            with st.spinner(f"Processing {sidebar_upload.name}..."):
                st.session_state.document_contents[sidebar_upload.name] = process_document(sidebar_upload)
            st.session_state.uploaded_files.append(file_details)
            st.sidebar.success(f"✅ {sidebar_upload.name} uploaded!")
            st.rerun()
        
        # Sidebar tips (Compact)
        st.markdown("---")
        st.markdown("**💡 Quick Tips:**")
        st.markdown("- Ask about studies or homework")
        st.markdown("- Upload files for context")
        st.markdown("- Be specific for better responses")
        
        # Show uploaded files (Styled)
        if st.session_state.uploaded_files:
            st.markdown("---")
            st.markdown("**📁 Your Files:**")
            for i, f in enumerate(st.session_state.uploaded_files):
                col1, col2 = st.columns([3, 1])
                with col1:
                    st.markdown(f"📄 {f['filename']}")
                with col2:
                    if st.button("🗑️", key=f"del_file_{i}", help="Remove file"):
                        fname = f['filename']
                        st.session_state.uploaded_files.pop(i)
                        st.session_state.document_contents.pop(fname, None)
                        st.rerun()
    
    # ---------------------------
    # Header (Exact Match to Quiz Structure: LectureBuddies - Chatbot)
    # ---------------------------
    st.markdown("<h1 class='main-title'>LectureBuddies</h1>", unsafe_allow_html=True)
    st.markdown("<h2 style='text-align: center; color: #4e54c8; font-weight: 600; font-family: Poppins, sans-serif;'>Lecturebuddies - Chatbot</h2>", unsafe_allow_html=True)
    st.markdown("<p class='tagline'>Your intelligent study companion—chat, summarize, and learn with AI ✨</p>", unsafe_allow_html=True)
    st.markdown("<hr class='gradient'>", unsafe_allow_html=True)
    
    # ---------------------------
    # Quick Actions (Styled to Match)
    # ---------------------------
    if not st.session_state.messages:
        st.markdown(
            """
            <div class="welcome-container">
                <h3 class="welcome-title">👋 Welcome to LectureBuddies Chat!</h3>
                <p class="welcome-text">Start chatting or try a quick action below to dive into your studies.</p>
            </div>
            """,
            unsafe_allow_html=True
        )
        
        col1, col2, col3 = st.columns(3)
        presets = {
            "📚 Help with Homework": "I need help understanding this homework assignment. Can you explain step-by-step?",
            "🔬 Explain a Concept": "I'm studying this concept but finding it difficult. Can you explain clearly with examples?",
            "💡 Study Tips": "I want to improve my study efficiency. What study strategies should I use?"
        }
        for col, (label, prompt) in zip([col1, col2, col3], presets.items()):
            with col:
                if st.button(label, key=label.replace(" ", "_").lower(), help="Start with this prompt"):
                    st.session_state.messages.append({"role": "user", "content": prompt})
                    with st.spinner("🤔 Thinking..."):
                        reply = get_groq_response(prompt)
                    st.session_state.messages.append({"role": "assistant", "content": reply})
                    st.rerun()
    
    # ---------------------------
    # Chat Display (Direct, No Container)
    # ---------------------------
    if st.session_state.messages:
        for msg in st.session_state.messages:
            if msg["role"] == "user":
                st.markdown(f"""
                <div class="user-message">
                    <div class="user-bubble">
                        {msg['content']}
                    </div>
                </div>
                """, unsafe_allow_html=True)
            else:
                st.markdown(f"""
                <div class="assistant-message">
                    <div class="assistant-bubble">
                        {msg['content']}
                    </div>
                </div>
                """, unsafe_allow_html=True)
    
    # ---------------------------
    # Chat Input (Professional Styling, Shorter Placeholder)
    # ---------------------------
    user_input = st.chat_input(placeholder="💬 Ask about studies or uploaded files...")
    if user_input and user_input.strip():
        # 🔑 Inject file content here
        final_input = inject_file_content(user_input.strip())
        
        st.session_state.messages.append({"role": "user", "content": user_input.strip()})
        with st.spinner("🤔 Thinking..."):
            reply = get_groq_response(final_input)
        st.session_state.messages.append({"role": "assistant", "content": reply})
        st.rerun()
    
    # ---------------------------
    # Footer (Compact)
    # ---------------------------
    st.markdown("---")
    st.markdown(
        "<p style='text-align: center; color: #888; font-size: 12px; font-family: Poppins, sans-serif; margin-bottom: 0;'>"
        "© 2023 LectureBuddies | Built with ❤️ for educational excellence | Powered by Groq AI</p>",
        unsafe_allow_html=True
    )

# ==========================
# QUIZ GENERATOR SECTION
# ==========================
def show_quiz_feature():
    """Display quiz generator feature with old design"""
    # Layout: Settings in Sidebar (like Chatbot), Main Content in Container
    # col_feature_sidebar, col_main_content = st.columns([1, 3]) # Removed to use genuine sidebar
    
    # Remove bottom spacing for this view
    st.markdown(
        """
        <style>
        .main .block-container { padding-bottom: 0 !important; margin-bottom: 0 !important; }
        </style>
        """,
        unsafe_allow_html=True
    )
    
    # LEFT: Feature-specific sidebar (Quiz Settings)
    with st.sidebar:
        st.markdown("### ⚙️ Quiz Settings")
        st.markdown("Customize your quiz generation.")
        
        st.markdown("**📊 Number of Questions**")
        num_questions = st.slider("How many questions?", min_value=1, max_value=20, value=st.session_state.num_questions, key="num_slider", label_visibility="collapsed")
        st.session_state.num_questions = num_questions

        st.markdown("**🎯 Difficulty Level**")
        difficulty = st.selectbox("Select difficulty:", ["Easy", "Medium", "Hard"], index=["Easy", "Medium", "Hard"].index(st.session_state.difficulty), key="diff_select", label_visibility="collapsed")
        st.session_state.difficulty = difficulty

        # Model and temperature controls
        model_options = [
            "llama-3.1-8b-instant",
            "llama-3.1-70b-versatile",
            "llama-3.2-11b-text-preview"
        ]
        st.session_state.quiz_model = st.selectbox("Model", model_options, index=model_options.index(st.session_state.quiz_model) if st.session_state.quiz_model in model_options else 0)
        st.session_state.quiz_temperature = st.slider("Creativity (temperature)", 0.0, 1.0, float(st.session_state.quiz_temperature), 0.1)
    
    # RIGHT: Main Quiz Interface
    with st.container():
        # Header
        st.markdown(
            """
            <div style="text-align: center; margin-bottom: 30px;">
                <h1 style="color: #4e54c8; font-size: 28px; font-weight: 700; margin-bottom: 10px; font-family: 'Poppins', sans-serif;">
                    lecturebuddies - Quiz Generator
                </h1>
                <p style="color: #666; font-size: 16px; font-family: 'Poppins', sans-serif;">
                    Transform your notes, lectures, or ideas into interactive quizzes with AI magic ✨
                </p>
                <hr style="border: 1px solid #4e54c8; margin: 20px 0;">
            </div>
            """,
            unsafe_allow_html=True
        )
        
        tab1, tab2, tab3 = st.tabs(["📁 Upload File", "💡 Enter Prompt", "📋 Paste Text"])

        # Tab 1 - File Upload
        with tab1:
            st.markdown("### Upload a document to generate a quiz from its content")
            uploaded_file = st.file_uploader(
                "Choose a file (PDF, DOCX, TXT)",
                type=["pdf", "docx", "txt"],
                help="Supported formats: PDF, Word documents, and plain text files."
            )

            if st.button("✨ Generate Quiz from File", key="file-btn", help=f"Create {num_questions} {difficulty} questions"):
                if uploaded_file:
                    content = extract_content_from_file(uploaded_file)
                    if content and "Error" not in content:
                        st.session_state.quiz_output = get_groq_quiz_response(content, num_questions, difficulty, model=st.session_state.quiz_model, temperature=st.session_state.quiz_temperature)
                    else:
                        st.error("Failed to extract content from file. Please try a different file or format.")
                else:
                    st.warning("Please upload a file first!")

        # Tab 2 - Prompt
        with tab2:
            st.markdown("### Describe a topic or provide content for the quiz")
            prompt_text = st.text_area(
                "Enter your topic, subject, or detailed content",
                placeholder="e.g., 'Explain photosynthesis and generate questions on it' or paste lecture notes...",
                height=100,
                help="Be as detailed as possible for better quizzes!"
            )

            if st.button("✨ Generate Quiz from Prompt", key="prompt-btn", help=f"Create {num_questions} {difficulty} questions"):
                if prompt_text.strip():
                    st.session_state.quiz_output = get_groq_quiz_response(prompt_text, num_questions, difficulty, model=st.session_state.quiz_model, temperature=st.session_state.quiz_temperature)
                else:
                    st.warning("Please enter some content or a topic!")

        # Tab 3 - Text Input
        with tab3:
            st.markdown("### Paste scanned or copied text directly")
            scanned_text = st.text_area(
                "Paste your text content here",
                placeholder="e.g., Copy-paste from a scanned PDF, image OCR, or notes...",
                height=100,
                help="Ideal for quick text from any source."
            )

            if st.button("✨ Generate Quiz from Text", key="scan-btn", help=f"Create {num_questions} {difficulty} questions"):
                if scanned_text.strip():
                    st.session_state.quiz_output = get_groq_quiz_response(scanned_text, num_questions, difficulty, model=st.session_state.quiz_model, temperature=st.session_state.quiz_temperature)
                else:
                    st.warning("Please paste some text!")

        # Display Quiz Results
        if st.session_state.quiz_output:
            if "Error" in st.session_state.quiz_output or "⚠️" in st.session_state.quiz_output or "❌" in st.session_state.quiz_output:
                st.error(st.session_state.quiz_output)
                if st.button("Clear", key="clear_error", help="Start over"):
                    st.session_state.quiz_output = None
                    st.rerun()
            else:
                st.markdown(
                    """
                    <div class="quiz-container">
                        <h3 style="color: #4e54c8; font-size: 22px; font-weight: 700; margin-bottom: 10px; text-align: center; font-family: 'Poppins', sans-serif;">
                            Your Generated Quiz
                        </h3>
                        <p style="text-align: center; color: #666; font-style: italic; font-size: 14px;">
                            {num_questions} questions at {difficulty} difficulty level
                        </p>
                    </div>
                    """.format(num_questions=st.session_state.num_questions, difficulty=st.session_state.difficulty),
                    unsafe_allow_html=True
                )

                st.markdown("### " + st.session_state.quiz_output)

                st.download_button(
                    label="Download Quiz as TXT",
                    data=st.session_state.quiz_output,
                    file_name=f"lecturebuddies_quiz_{st.session_state.num_questions}q_{st.session_state.difficulty.lower()}.txt",
                    mime="text/plain",
                    help="Save your quiz for later use!"
                )

                if st.button("Generate New Quiz", key="clear", help="Clear and start over"):
                    st.session_state.quiz_output = None
                    st.rerun()

# ==========================
# LIVE LECTURE RECORDING SECTION
# ==========================
def show_recording_feature():
    """Display simplified speech-to-text transcription via file upload or Real-Time"""
    
    # 1. Main heading
    st.markdown(
        """
        <div style="text-align: center; margin-bottom: 30px;">
            <h1 style="color: #4e54c8; font-size: 32px; font-weight: 700; margin-bottom: 10px; font-family: 'Poppins', sans-serif;">
                🎤 LectureBuddies — Speech to Text
            </h1>
            <p style="color: #666; font-size: 16px; font-family: 'Poppins', sans-serif;">
                Upload voice recordings or record live audio for real-time transcription
            </p>
        </div>
        """,
        unsafe_allow_html=True
    )

    SAVE_DIR = "temp_recordings"
    os.makedirs(SAVE_DIR, exist_ok=True)

    # 2. Load Whisper Model (Updated to 'tiny.en' for speed as requested)
    @st.cache_resource
    def load_whisper_model_simple(model_size="tiny.en", device="cpu", compute_type="int8"):
        # Try to use GPU if available, else CPU
        try:
            return WhisperModel(model_size, device="cuda", compute_type="int8_float16")
        except:
            return WhisperModel(model_size, device="cpu", compute_type="int8")

    with st.spinner(f"Loading High-Speed Model..."):
        model = load_whisper_model_simple()

    # 3. Tabs for different modes
    tab1, tab2 = st.tabs(["📁 Upload Audio Files", "🎤 Realtime Transcript"])

    # ==========================
    # TAB 1: Upload Audio Files (Existing Working Code)
    # ==========================
    with tab1:
        st.markdown("### Upload audio files for transcription")
        
        uploaded_files = st.file_uploader(
            "Upload audio file(s)",
            type=["wav", "mp3", "m4a", "flac", "ogg"],
            help="Supported formats: WAV, MP3, M4A, FLAC, OGG",
            accept_multiple_files=True,
            key="upload_files"
        )

        if uploaded_files:
            if st.button("🎯 Transcribe Files", key="transcribe_all"):
                for idx, uploaded_file in enumerate(uploaded_files, start=1):
                    with st.spinner(f"Transcribing {uploaded_file.name} ({idx}/{len(uploaded_files)})..."):
                        # Save to temp
                        temp_path = os.path.join(SAVE_DIR, f"uploaded_{int(time.time())}_{uploaded_file.name}")
                        with open(temp_path, "wb") as f:
                            f.write(uploaded_file.read())

                        try:
                            segments, info = model.transcribe(temp_path, beam_size=5)
                            transcription_text = "".join([seg.text for seg in segments])

                            st.success(f"✅ Transcription completed: {uploaded_file.name}")

                            # Display and downloads
                            with st.expander(f"📝 {uploaded_file.name} — View Transcription", expanded=True):
                                st.text_area("Result", value=transcription_text, height=200, key=f"tx_{idx}")
                                st.download_button(
                                    "📥 Download Transcription",
                                    transcription_text,
                                    file_name=f"transcript_{uploaded_file.name}.txt"
                                )
                        except Exception as e:
                            st.error(f"❌ Error transcribing {uploaded_file.name}: {str(e)}")
                        finally:
                            if os.path.exists(temp_path):
                                os.remove(temp_path)

    # ==========================
    # TAB 2: Realtime Transcript (UPDATED FOR INTERNET)
    # ==========================
    with tab2:
        st.markdown("### ⚡ Fast Real-time Transcription")
        
        # Initialize session state for transcript storage
        if "live_transcript" not in st.session_state:
            st.session_state.live_transcript = ""

        col1, col2 = st.columns([1, 2])

        with col1:
            st.info("💡 **Instructions:**\n1. Click 'Record' below.\n2. Speak.\n3. Click 'Stop' to send audio to AI.\n(Note: On the internet, audio processes after you stop speaking).")
            
            # Use web-recorder instead of local mic
            audio_data = mic_recorder(
                start_prompt="🔴 Start Recording",
                stop_prompt="⏹️ Stop & Transcribe",
                key='recorder',
                format='wav',
                use_container_width=True
            )
            
            # Download Button (Visible if we have text)
            if st.session_state.live_transcript:
                st.download_button(
                    label="📥 Download Transcript",
                    data=st.session_state.live_transcript,
                    file_name="live_transcript.txt",
                    mime="text/plain",
                    use_container_width=True
                )
                
            if st.button("🗑️ Clear Transcript", use_container_width=True):
                st.session_state.live_transcript = ""
                st.rerun()

        with col2:
            st.markdown("#### 📝 Live Output")
            
            # Logic to process audio when user stops recording
            if audio_data is not None:
                with st.spinner("⚡ Transcribing audio chunk..."):
                    # Save bytes to temp file for Whisper
                    with tempfile.NamedTemporaryFile(suffix=".wav", delete=False) as tmp_audio:
                        tmp_audio.write(audio_data['bytes'])
                        tmp_audio_path = tmp_audio.name
                    
                    try:
                        # Transcribe
                        segments, _ = model.transcribe(tmp_audio_path, beam_size=1, language="en")
                        new_text = "".join([s.text for s in segments]).strip()
                        
                        if new_text:
                            st.session_state.live_transcript += new_text + " "
                            st.success("Audio processed!")
                    except Exception as e:
                        st.error(f"Error processing audio: {e}")
                    finally:
                        if os.path.exists(tmp_audio_path):
                            os.remove(tmp_audio_path)

            # Show existing transcript
            st.text_area("Transcript", value=st.session_state.live_transcript, height=400, disabled=True)

# ==========================
# FLASH CARDS FEATURE
# ==========================
def show_flashcards_feature():
    """Display flash cards generator feature"""
    st.markdown(
        """
        <div style="text-align: center; margin-bottom: 30px;">
            <h1 style="color: #4e54c8; font-size: 32px; font-weight: 700; margin-bottom: 10px; font-family: 'Poppins', sans-serif;">
                📖 Flash Cards Generator
            </h1>
            <p style="color: #666; font-size: 16px; font-family: 'Poppins', sans-serif;">
                Create interactive flash cards from your study materials
            </p>
        </div>
        """,
        unsafe_allow_html=True
    )
    
    col1, col2 = st.columns([1, 2])
    
    with col1:
        st.markdown("### ⚙️ Settings")
        st.markdown("---")
        
        # Flash card settings
        num_cards = st.slider("Number of Cards", 1, 50, 10)
        difficulty = st.selectbox("Difficulty", ["Easy", "Medium", "Hard"])
        subject = st.text_input("Subject/Topic", placeholder="e.g., Biology, History")
        
        st.markdown("---")
        st.markdown("**📄 Upload Content**")
        uploaded_file = st.file_uploader(
            "Choose a file",
            type=['txt', 'pdf', 'docx'],
            help="Upload study material to generate flash cards"
        )
    
    with col2:
        st.markdown("### 📝 Content Input")
        
        content_input = st.text_area(
            "Or paste your study content here",
            placeholder="Paste your study material, lecture notes, or textbook content...",
            height=200
        )
        
        if st.button("🎯 Generate Flash Cards", key="generate_cards", use_container_width=True):
            if content_input or uploaded_file:
                with st.spinner("Generating flash cards..."):
                    # Generate flash cards using AI
                    flashcards = generate_flashcards(content_input or "Sample content", num_cards, difficulty, subject)
                    st.session_state.flashcards = flashcards
                    st.success(f"✅ Generated {len(flashcards)} flash cards!")
            else:
                st.warning("Please provide content to generate flash cards")
        
        # Display generated flash cards
        if st.session_state.get('flashcards'):
            st.markdown("### 🃏 Your Flash Cards")
            display_flashcards(st.session_state.flashcards)

def generate_flashcards(content, num_cards, difficulty, subject):
    """Generate flash cards from content using AI"""
    if not api_key:
        return [{"front": "Error: No API key", "back": "Please set GROQ_API_KEY in your .env file"}]
    
    url = "https://api.groq.com/openai/v1/chat/completions"
    headers = {"Authorization": f"Bearer {api_key}", "Content-Type": "application/json"}
    
    prompt = f"""
    Create {num_cards} flash cards from this content for {subject or 'general study'}.
    Difficulty level: {difficulty}
    
    Content: {content[:2000]}
    
    Format each card as:
    Front: [Question or term]
    Back: [Answer or definition]
    
    Make them educational and useful for studying.
    """
    
    messages = [
        {"role": "system", "content": "You are a helpful study assistant that creates educational flash cards."},
        {"role": "user", "content": prompt}
    ]
    
    payload = {"model": "llama-3.1-8b-instant", "messages": messages, "temperature": 0.7, "max_tokens": 1000}
    
    try:
        response = requests.post(url, headers=headers, json=payload, timeout=30)
        if response.status_code == 200:
            data = response.json()
            content = data.get("choices", [{}])[0].get("message", {}).get("content", "")
            return parse_flashcards(content)
        else:
            return [{"front": "Error generating cards", "back": f"API Error: {response.status_code}"}]
    except Exception as e:
        return [{"front": "Error", "back": f"Failed to generate cards: {str(e)}"}]

def parse_flashcards(content):
    """Parse AI response into flash card format"""
    cards = []
    lines = content.split('\n')
    current_card = {}
    
    for line in lines:
        line = line.strip()
        if line.startswith('Front:'):
            if current_card:
                cards.append(current_card)
            current_card = {"front": line.replace('Front:', '').strip(), "back": ""}
        elif line.startswith('Back:'):
            current_card["back"] = line.replace('Back:', '').strip()
    
    if current_card:
        cards.append(current_card)
    
    return cards if cards else [{"front": "Sample Question", "back": "Sample Answer"}]

def display_flashcards(cards):
    """Display flash cards in an interactive format"""
    if not cards:
        return
    
    # Initialize session state for current card
    if 'current_card_index' not in st.session_state:
        st.session_state.current_card_index = 0
    
    current_index = st.session_state.current_card_index
    current_card = cards[current_index]
    
    # Card display
    st.markdown(
        f"""
        <div style="background: white; padding: 30px; border-radius: 15px; box-shadow: 0 4px 15px rgba(0,0,0,0.1); text-align: center; margin: 20px 0;">
            <h3 style="color: #4e54c8; margin-bottom: 20px;">Card {current_index + 1} of {len(cards)}</h3>
            <div style="background: #f8f9ff; padding: 20px; border-radius: 10px; margin: 20px 0;">
                <h4 style="color: #2c3e50; margin-bottom: 15px;">Front:</h4>
                <p style="font-size: 18px; color: #333;">{current_card['front']}</p>
            </div>
        </div>
        """,
        unsafe_allow_html=True
    )
    
    # Show answer button
    if st.button("👁️ Show Answer", key="show_answer"):
        st.markdown(
            f"""
            <div style="background: #e8f5e8; padding: 20px; border-radius: 10px; margin: 20px 0;">
                <h4 style="color: #2c3e50; margin-bottom: 15px;">Back:</h4>
                <p style="font-size: 18px; color: #333;">{current_card['back']}</p>
            </div>
            """,
            unsafe_allow_html=True
        )
    
    # Navigation
    col1, col2, col3 = st.columns([1, 2, 1])
    
    with col1:
        if st.button("⬅️ Previous", key="prev_card", disabled=current_index == 0):
            st.session_state.current_card_index = max(0, current_index - 1)
            st.rerun()
    
    with col2:
        st.markdown(f"<div style='text-align: center; color: #666;'>Card {current_index + 1} of {len(cards)}</div>", unsafe_allow_html=True)
    
    with col3:
        if st.button("Next ➡️", key="next_card", disabled=current_index == len(cards) - 1):
            st.session_state.current_card_index = min(len(cards) - 1, current_index + 1)
            st.rerun()

# ==========================
# MULTILINGUAL TRANSLATION FEATURE
# ==========================
def show_translation_feature():
    """Display multilingual translation feature"""
    st.markdown(
        """
        <div style="text-align: center; margin-bottom: 30px;">
            <h1 style="color: #4e54c8; font-size: 32px; font-weight: 700; margin-bottom: 10px; font-family: 'Poppins', sans-serif;">
                🌐 Multilingual Translation
            </h1>
            <p style="color: #666; font-size: 16px; font-family: 'Poppins', sans-serif;">
                Translate text to your selected language (source auto-detected)
            </p>
        </div>
        """,
        unsafe_allow_html=True
    )
    
    col1, col2 = st.columns([1, 2])
    
    with col1:
        st.markdown("### ⚙️ Translation Settings")
        st.markdown("---")
        
        # Language selection (target only)
        languages = {
            "English": "en",
            "Spanish": "es", 
            "French": "fr",
            "German": "de",
            "Italian": "it",
            "Portuguese": "pt",
            "Chinese": "zh",
            "Japanese": "ja",
            "Korean": "ko",
            "Arabic": "ar",
            "Hindi": "hi",
            "Russian": "ru",
            "Urdu": "ur",
            "Punjabi": "pa",
            "Sindhi": "sd"
        }
        target_lang_label = st.selectbox("To Language", list(languages.keys()))
        
        st.markdown("---")
        st.markdown("**📄 Upload Document**")
        uploaded_file = st.file_uploader(
            "Choose a text file",
            type=['txt', 'docx'],
            help="Upload a document to translate"
        )
    
    with col2:
        st.markdown("### 📝 Text Translation")
        
        # Text input
        text_input = st.text_area(
            "Enter text to translate",
            placeholder="Type or paste your text here...",
            height=200
        )
        
        if st.button("🔄 Translate", key="translate_text", use_container_width=True):
            if text_input or uploaded_file:
                with st.spinner("Translating..."):
                    if uploaded_file:
                        # Process uploaded file
                        content = extract_content_from_file(uploaded_file)
                        translated = translate_text(content, target_lang_label)
                    else:
                        translated = translate_text(text_input, target_lang_label)
                    
                    st.session_state.translation_result = translated
                    st.success("✅ Translation completed!")
            else:
                st.warning("Please provide text to translate")
        
        # Display translation result
        if st.session_state.get('translation_result'):
            st.markdown("### 🌐 Translation Result")
            st.text_area("Translated Text", value=st.session_state.translation_result, height=200, key="translation_display")
            
            # Download button
            st.download_button(
                "📥 Download Translation",
                st.session_state.translation_result,
                file_name=f"translation_to_{languages[target_lang_label]}.txt",
                mime="text/plain"
            )

def translate_text(text, target_language_label):
    """Translate text using AI; auto-detect source language"""
    if not api_key:
        return "Error: No API key available"
    
    url = "https://api.groq.com/openai/v1/chat/completions"
    headers = {"Authorization": f"Bearer {api_key}", "Content-Type": "application/json"}
    
    prompt = (
        f"Detect the source language and translate the following text to {target_language_label}. "
        f"Only return the translated text, no explanations or prefixes.\n\n{text}"
    )
    
    messages = [
        {"role": "system", "content": "You are a professional translator. Translate accurately and naturally."},
        {"role": "user", "content": prompt}
    ]
    
    payload = {"model": "llama-3.1-8b-instant", "messages": messages, "temperature": 0.3, "max_tokens": 1000}
    
    try:
        response = requests.post(url, headers=headers, json=payload, timeout=30)
        if response.status_code == 200:
            data = response.json()
            return data.get("choices", [{}])[0].get("message", {}).get("content", "Translation failed")
        else:
            return f"Translation error: {response.status_code}"
    except Exception as e:
        return f"Translation failed: {str(e)}"

# ==========================
# NOTES MANAGER FEATURE
# ==========================
def show_notes_feature():
    """Display notes manager with import/export and organization"""
    st.markdown(
        """
        <div style="text-align: center; margin-bottom: 30px;">
            <h1 style="color: #4e54c8; font-size: 32px; font-weight: 700; margin-bottom: 10px; font-family: 'Poppins', sans-serif;">
                📋 Notes Manager
            </h1>
            <p style="color: #666; font-size: 16px; font-family: 'Poppins', sans-serif;">
                Organize, import, and export your study notes with rich text formatting
            </p>
        </div>
        """,
        unsafe_allow_html=True
    )
    
    # Initialize notes in session state
    if 'notes' not in st.session_state:
        st.session_state.notes = []
    
    col1, col2 = st.columns([1, 2])
    
    with col1:
        st.markdown("### 📁 Note Organization")
        st.markdown("---")
        
        # Categories
        categories = ["General", "Mathematics", "Science", "History", "Language", "Other"]
        selected_category = st.selectbox("Category", categories)
        
        # Tags
        tags_input = st.text_input("Tags (comma-separated)", placeholder="e.g., important, exam, review")
        
        st.markdown("---")
        st.markdown("**📤 Import/Export**")
        
        # Import notes
        uploaded_notes = st.file_uploader(
            "Import Notes",
            type=['txt', 'json'],
            help="Import notes from file"
        )
        
        if uploaded_notes:
            if st.button("📥 Import Notes", key="import_notes"):
                try:
                    if uploaded_notes.type == "application/json":
                        notes_data = json.loads(uploaded_notes.read())
                        st.session_state.notes.extend(notes_data)
                    else:
                        content = uploaded_notes.read().decode('utf-8')
                        new_note = {
                            "title": uploaded_notes.name,
                            "content": content,
                            "category": selected_category,
                            "tags": tags_input.split(',') if tags_input else [],
                            "created": time.time()
                        }
                        st.session_state.notes.append(new_note)
                    st.success("✅ Notes imported successfully!")
                except Exception as e:
                    st.error(f"❌ Import failed: {str(e)}")
        
        # Export notes
        if st.session_state.notes:
            if st.button("📤 Export All Notes", key="export_notes"):
                notes_json = json.dumps(st.session_state.notes, indent=2)
                st.download_button(
                    "💾 Download Notes (JSON)",
                    notes_json,
                    file_name=f"lecturebuddies_notes_{int(time.time())}.json",
                    mime="application/json"
                )
    
    with col2:
        st.markdown("### ✏️ Create New Note")
        
        # Note creation form
        note_title = st.text_input("Note Title", placeholder="Enter note title")
        
        # Rich text editor (simplified)
        note_content = st.text_area(
            "Note Content",
            placeholder="Write your note here...\n\nYou can use basic formatting:\n**Bold text**\n*Italic text*\n# Heading\n- Bullet point",
            height=300
        )
        
        if st.button("💾 Save Note", key="save_note", use_container_width=True):
            if note_title and note_content:
                new_note = {
                    "title": note_title,
                    "content": note_content,
                    "category": selected_category,
                    "tags": [tag.strip() for tag in tags_input.split(',')] if tags_input else [],
                    "created": time.time(),
                    "modified": time.time()
                }
                st.session_state.notes.append(new_note)
                st.success("✅ Note saved successfully!")
                st.rerun()
            else:
                st.warning("Please enter both title and content")
        
        # Display existing notes
        if st.session_state.notes:
            st.markdown("### 📚 Your Notes")
            for i, note in enumerate(st.session_state.notes):
                with st.expander(f"📝 {note['title']} ({note['category']})"):
                    st.markdown(f"**Created:** {time.ctime(note['created'])}")
                    st.markdown(f"**Tags:** {', '.join(note['tags']) if note['tags'] else 'None'}")
                    st.markdown("**Content:**")
                    st.markdown(note['content'])
                    
                    col_del, col_edit = st.columns(2)
                    with col_del:
                        if st.button("🗑️ Delete", key=f"del_note_{i}"):
                            st.session_state.notes.pop(i)
                            st.rerun()
                    with col_edit:
                        if st.button("✏️ Edit", key=f"edit_note_{i}"):
                            st.session_state.editing_note = i
                            st.rerun()

# ==========================
# ADMIN DASHBOARD FEATURE
# ==========================
def show_admin_feature():
    """Display admin dashboard for system management"""
    st.markdown(
        """
        <div style="text-align: center; margin-bottom: 30px;">
            <h1 style="color: #4e54c8; font-size: 32px; font-weight: 700; margin-bottom: 10px; font-family: 'Poppins', sans-serif;">
                👨‍💼 Admin Dashboard
            </h1>
            <p style="color: #666; font-size: 16px; font-family: 'Poppins', sans-serif;">
                System management and analytics for administrators
            </p>
        </div>
        """,
        unsafe_allow_html=True
    )
    
    # Check if user is admin
    if st.session_state.current_user != "admin":
        st.warning("🔒 Admin access required. Please login as admin.")
        return
    
    col1, col2, col3 = st.columns(3)
    
    with col1:
        st.markdown("### 📊 System Stats")
        st.metric("Total Users", "1,234")
        st.metric("Active Sessions", "45")
        st.metric("Storage Used", "2.3 GB")
    
    with col2:
        st.markdown("### 🎯 Feature Usage")
        st.metric("Chatbot Queries", "5,678")
        st.metric("Quiz Generated", "234")
        st.metric("Recordings Made", "89")
    
    with col3:
        st.markdown("### ⚙️ System Health")
        st.metric("API Status", "✅ Online")
        st.metric("Database", "✅ Connected")
        st.metric("Storage", "✅ Available")
    
    st.markdown("---")
    
    # System management
    st.markdown("### 🔧 System Management")
    
    tab1, tab2, tab3 = st.tabs(["Users", "Settings", "Logs"])
    
    with tab1:
        st.markdown("#### 👥 User Management")
        st.dataframe({
            "Username": ["admin", "student1", "student2"],
            "Role": ["Admin", "Student", "Student"],
            "Last Login": ["2024-01-15", "2024-01-14", "2024-01-13"],
            "Status": ["Active", "Active", "Inactive"]
        })
    
    with tab2:
        st.markdown("#### ⚙️ System Settings")
        st.checkbox("Enable Registration", value=True)
        st.checkbox("Maintenance Mode", value=False)
        st.slider("Max File Size (MB)", 1, 100, 10)
    
    with tab3:
        st.markdown("#### 📋 System Logs")
        st.text_area("Recent Logs", value="2024-01-15 10:30:15 - User login: admin\n2024-01-15 10:25:32 - Quiz generated: student1\n2024-01-15 10:20:45 - Recording started: student2", height=200)

# ==========================
# CATEGORIZED SEARCH FEATURE
# ==========================
def show_search_feature():
    """Display categorized search functionality"""
    st.markdown(
        """
        <div style="text-align: center; margin-bottom: 30px;">
            <h1 style="color: #4e54c8; font-size: 32px; font-weight: 700; margin-bottom: 10px; font-family: 'Poppins', sans-serif;">
                🔍 Categorized Search
            </h1>
            <p style="color: #666; font-size: 16px; font-family: 'Poppins', sans-serif;">
                Search across all your content with advanced filtering and categorization
            </p>
        </div>
        """,
        unsafe_allow_html=True
    )
    
    col1, col2 = st.columns([1, 2])
    
    with col1:
        st.markdown("### 🔍 Search Filters")
        st.markdown("---")
        
        # Search categories
        search_categories = st.multiselect(
            "Search in:",
            ["Chat History", "Notes", "Quizzes", "Recordings", "Flash Cards"],
            default=["Chat History", "Notes"]
        )
        
        # Date range
        date_range = st.date_input("Date Range", value=[])
        
        # Content type
        content_type = st.selectbox("Content Type", ["All", "Text", "Audio", "Images", "Documents"])
        
        # Tags filter
        tags_filter = st.text_input("Tags", placeholder="e.g., important, exam")
    
    with col2:
        st.markdown("### 🔎 Search Query")
        
        search_query = st.text_input(
            "Enter your search query",
            placeholder="Search for specific topics, keywords, or phrases...",
            key="search_input"
        )
        
        if st.button("🔍 Search", key="perform_search", use_container_width=True):
            if search_query:
                with st.spinner("Searching..."):
                    results = perform_search(search_query, search_categories, date_range, content_type, tags_filter)
                    st.session_state.search_results = results
                    st.success(f"✅ Found {len(results)} results!")
            else:
                st.warning("Please enter a search query")
        
        # Display search results
        if st.session_state.get('search_results'):
            st.markdown("### 📋 Search Results")
            display_search_results(st.session_state.search_results)

def perform_search(query, categories, date_range, content_type, tags_filter):
    """Perform categorized search"""
    # This is a simplified search implementation
    # In a real application, you would search through actual data
    results = []
    
    # Simulate search results
    sample_results = [
        {
            "title": "Chemistry Notes - Organic Compounds",
            "content": "Organic compounds are carbon-based molecules...",
            "category": "Notes",
            "date": "2024-01-15",
            "tags": ["chemistry", "organic", "important"],
            "relevance": 0.95
        },
        {
            "title": "Math Quiz - Calculus",
            "content": "Quiz questions about derivatives and integrals...",
            "category": "Quizzes", 
            "date": "2024-01-14",
            "tags": ["math", "calculus", "quiz"],
            "relevance": 0.87
        },
        {
            "title": "Biology Lecture Recording",
            "content": "Transcription of biology lecture about cell division...",
            "category": "Recordings",
            "date": "2024-01-13",
            "tags": ["biology", "lecture", "cell"],
            "relevance": 0.82
        }
    ]
    
    # Filter results based on search criteria
    for result in sample_results:
        if result['category'] in categories:
            if query.lower() in result['title'].lower() or query.lower() in result['content'].lower():
                results.append(result)
    
    return sorted(results, key=lambda x: x['relevance'], reverse=True)

def display_search_results(results):
    """Display search results in a formatted way"""
    for i, result in enumerate(results):
        with st.expander(f"📄 {result['title']} (Relevance: {result['relevance']:.2f})"):
            st.markdown(f"**Category:** {result['category']}")
            st.markdown(f"**Date:** {result['date']}")
            st.markdown(f"**Tags:** {', '.join(result['tags'])}")
            st.markdown(f"**Content:** {result['content'][:200]}...")
            
            if st.button("📖 View Full", key=f"view_result_{i}"):
                st.markdown(f"**Full Content:**\n\n{result['content']}")

# ==========================
# OFFLINE MODE FEATURE
# ==========================
def show_offline_feature():
    """Display offline mode functionality"""
    st.markdown(
        """
        <div style="text-align: center; margin-bottom: 30px;">
            <h1 style="color: #4e54c8; font-size: 32px; font-weight: 700; margin-bottom: 10px; font-family: 'Poppins', sans-serif;">
                📱 Offline Mode
            </h1>
            <p style="color: #666; font-size: 16px; font-family: 'Poppins', sans-serif;">
                Access your content and basic features without internet connection
            </p>
        </div>
        """,
        unsafe_allow_html=True
    )
    
    # Check internet connection
    try:
        requests.get("https://www.google.com", timeout=5)
        internet_status = "🟢 Online"
        connection_color = "#28a745"
    except:
        internet_status = "🔴 Offline"
        connection_color = "#dc3545"
    
    st.markdown(
        f"""
        <div style="background: #f8f9fa; padding: 20px; border-radius: 10px; margin: 20px 0; text-align: center;">
            <h3 style="color: {connection_color}; margin-bottom: 10px;">Connection Status: {internet_status}</h3>
            <p style="color: #666;">Current mode: {'Offline' if 'Offline' in internet_status else 'Online'}</p>
        </div>
        """,
        unsafe_allow_html=True
    )
    
    col1, col2 = st.columns(2)
    
    with col1:
        st.markdown("### 📱 Offline Features")
        st.markdown("---")
        
        offline_features = [
            "📚 View saved notes",
            "🃏 Review flash cards", 
            "📝 Read saved quizzes",
            "🎧 Play downloaded recordings",
            "🔍 Search local content",
            "📊 View offline analytics"
        ]
        
        for feature in offline_features:
            st.markdown(f"✅ {feature}")
        
        st.markdown("---")
        st.markdown("### 💾 Sync Options")
        
        if st.button("🔄 Sync Now", key="sync_now"):
            if "Offline" in internet_status:
                st.warning("⚠️ No internet connection. Sync will happen when you're back online.")
            else:
                st.success("✅ Sync completed! All offline content updated.")
    
    with col2:
        st.markdown("### 📊 Offline Content")
        
        # Display offline content stats
        st.metric("Saved Notes", "23")
        st.metric("Flash Cards", "156")
        st.metric("Quizzes", "12")
        st.metric("Recordings", "8")
        
        st.markdown("---")
        st.markdown("### ⚙️ Offline Settings")
        
        auto_sync = st.checkbox("Auto-sync when online", value=True)
        download_quality = st.selectbox("Download Quality", ["High", "Medium", "Low"])
        storage_limit = st.slider("Storage Limit (GB)", 1, 10, 5)
        
        if st.button("💾 Download for Offline", key="download_offline"):
            st.info("📥 Downloading content for offline use...")
            st.progress(0.7)
            st.success("✅ Offline content ready!")

# ==========================
# HELPER FUNCTIONS
# ==========================
def inject_file_content(user_message: str) -> str:
    """Replace file references in user message with extracted text"""
    for fname, content in st.session_state.document_contents.items():
        if fname.lower() in user_message.lower():
            extracted = content if content.strip() else "[No text extracted from this file]"
            user_message = user_message.replace(
                fname,
                f"(Extracted content: {extracted[:1000]}...)"
            )
    return user_message

def get_groq_response(user_input, model="llama-3.1-8b-instant", temperature=0.7):
    """Send query + context to Groq API and return assistant response"""
    if not api_key or api_key.strip() == "":
        return "Missing API key. Please set GROQ_API_KEY in your .env file."

    url = "https://api.groq.com/openai/v1/chat/completions"
    headers = {"Authorization": f"Bearer {api_key}", "Content-Type": "application/json"}

    # Build document context
    doc_context = ""
    if st.session_state.document_contents:
        doc_context = "\n\n**Available Documents:**\n"
        for fname, content in st.session_state.document_contents.items():
            snippet = content[:1000] if content else "[No extractable text]"
            doc_context += f"\n--- {fname} ---\n{snippet}...\n"

    # Build conversation
    messages = [{"role": m["role"], "content": m["content"]} for m in st.session_state.messages]
    system_msg = (
        f"You are LectureBuddies, an AI chatbot designed for education. "
        f"Answer clearly, summarize effectively, and explain concepts step by step."
        f"{' Available Documents: ' + doc_context if doc_context else ''}\n\n"
        "Guidelines:\n"
        "Education-focused\n"
        "Summarization expert\n"
        "Clarity first (simple language, then details)\n"
        "Confidence + accuracy\n"
        "Break down topics step-by-step, use examples, and stay professional yet supportive."
    )
    messages.insert(0, {"role": "system", "content": system_msg})
    messages.append({"role": "user", "content": user_input})

    payload = {"model": model, "messages": messages, "temperature": float(temperature), "max_tokens": 1000}

    try:
        resp = requests.post(url, headers=headers, json=payload, timeout=30)
        if resp.status_code == 200:
            data = resp.json()
            return data.get("choices", [{}])[0].get("message", {}).get("content", "No response received.")
        elif resp.status_code == 401:
            return "Invalid API key. Please check GROQ_API_KEY in your .env file."
        elif resp.status_code == 429:
            return "Too many requests. Please slow down and retry shortly."
        else:
            return f"API Error {resp.status_code}: {resp.text}"
    except requests.exceptions.Timeout:
        return "Request timed out. Please retry."
    except requests.exceptions.RequestException as e:
        return f"Network error: {e}"
    except Exception as e:
        return f"Unexpected error: {e}"

def process_document(uploaded_file, doc_processor):
    """Extract text from uploaded documents"""
    try:
        temp_path = os.path.join("temp", uploaded_file.name)
        os.makedirs("temp", exist_ok=True)
        with open(temp_path, "wb") as f:
            f.write(uploaded_file.read())

        content = doc_processor.process_document(temp_path)
        return content if content.strip() else "[No text extracted]"
    except Exception as e:
        return f"[File processing error: {e}]"

def get_groq_quiz_response(content, num_questions=5, difficulty="Medium", model="llama-3.1-8b-instant", temperature=0.7):
    """Send content to Groq API and get quiz questions back"""
    url = "https://api.groq.com/openai/v1/chat/completions"
    headers = {"Authorization": f"Bearer {api_key}", "Content-Type": "application/json"}

    difficulty_map = {"Easy": "simple and straightforward", "Medium": "balanced and informative", "Hard": "challenging and detailed"}
    diff_desc = difficulty_map.get(difficulty, "balanced")

    system_msg = (
        "You are LectureBuddies Quiz Generator. "
        "Your task is to generate high-quality multiple-choice quizzes (MCQs) from educational material. "
        f"Generate exactly {num_questions} MCQs. Each question should have 4 options (A, B, C, D), one correct answer, "
        "and clearly mark the correct answer (e.g., Correct: A). "
        f"Make questions {diff_desc} in difficulty. "
        "Format the output neatly with numbered questions, bold question text, and labeled options. "
        "End with a summary of correct answers."
    )

    messages = [
        {"role": "system", "content": system_msg},
        {"role": "user", "content": f"Generate a quiz from this content:\n\n{content}"}
    ]

    payload = {"model": model, "messages": messages, "temperature": float(temperature), "max_tokens": 1200}

    try:
        with st.spinner(f"Generating {num_questions} {difficulty} quiz questions..."):
            resp = requests.post(url, headers=headers, json=payload, timeout=30)
        if resp.status_code == 200:
            data = resp.json()
            return data.get("choices", [{}])[0].get("message", {}).get("content", "No response received.")
        elif resp.status_code == 401:
            return "Invalid API key. Please check GROQ_API_KEY in your .env file."
        elif resp.status_code == 429:
            return "Too many requests. Please retry shortly."
        else:
            return f"API Error {resp.status_code}: {resp.text}"
    except requests.exceptions.Timeout:
        return "Request timed out. Please retry."
    except Exception as e:
        return f"Error: {str(e)}"

def extract_content_from_file(uploaded_file):
    """Extract text content from uploaded file (PDF, DOCX, TXT)"""
    file_type = uploaded_file.type if hasattr(uploaded_file, 'type') else uploaded_file.name.split('.')[-1].lower()

    try:
        if file_type == "application/pdf" or file_type.endswith('.pdf'):
            pdf_reader = PyPDF2.PdfReader(io.BytesIO(uploaded_file.read()))
            content = ""
            for page in pdf_reader.pages:
                text = page.extract_text() or ""
                content += text + "\n"
            return content.strip()
        elif file_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document" or file_type.endswith('.docx'):
            doc = Document(io.BytesIO(uploaded_file.read()))
            content = "\n".join([para.text for para in doc.paragraphs if para.text.strip()])
            return content.strip()
        else:  # TXT or other text-based
            return uploaded_file.read().decode("utf-8", errors="ignore").strip()
    except Exception as e:
        return f"Error extracting content: {str(e)}"

# ==========================
# MAIN APPLICATION LOGIC
# ==========================
def main():
    """Main application entry point"""
    if not st.session_state.authenticated:
        show_login_page()
    else:
        show_dashboard()

if __name__ == "__main__":
    main()
